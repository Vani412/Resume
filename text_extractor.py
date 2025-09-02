import PyPDF2
import docx
from typing import Optional
import streamlit as st

class TextExtractor:
    """
    Handles text extraction from PDF and DOCX files
    """
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """
        Extract text from uploaded resume file
        
        Args:
            file_path (str): Path to the uploaded file
            file_type (str): MIME type of the file
            
        Returns:
            str: Extracted text content
        """
        try:
            if 'pdf' in file_type.lower():
                return self._extract_from_pdf(file_path)
            elif 'docx' in file_type.lower() or 'document' in file_type.lower():
                return self._extract_from_docx(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            st.error(f"Error extracting text: {str(e)}")
            raise
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from all pages
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    text += page_text + "\n"
                    
        except Exception as e:
            # Try alternative PDF extraction method
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except ImportError:
                raise Exception(f"PDF extraction failed: {str(e)}. Please ensure the PDF is not password-protected or corrupted.")
            except Exception as e2:
                raise Exception(f"PDF extraction failed with both methods: {str(e)} | {str(e2)}")
        
        return self._clean_extracted_text(text)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        text = ""
        
        try:
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
        except Exception as e:
            raise Exception(f"DOCX extraction failed: {str(e)}. Please ensure the file is not corrupted.")
        
        return self._clean_extracted_text(text)
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        import re
        text = re.sub(r'\n\s*\n', '\n\n', text)  # Multiple newlines to double newline
        text = re.sub(r'\s+', ' ', text)  # Multiple spaces to single space
        text = text.strip()
        
        return text
    
    def validate_file_type(self, file_type: str) -> bool:
        """Validate if the file type is supported"""
        supported_types = [
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]
        
        return any(supported_type in file_type for supported_type in supported_types)
    
    def get_file_info(self, uploaded_file) -> dict:
        """Get information about the uploaded file"""
        return {
            'name': uploaded_file.name,
            'size': uploaded_file.size,
            'type': uploaded_file.type,
            'size_mb': round(uploaded_file.size / (1024 * 1024), 2)
        }
